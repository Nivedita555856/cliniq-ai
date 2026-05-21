# streamlit_run.py - Medical Report Analyzer
# Complete working version with Groq Llama 3 + RAG Integration
# Supports: CBC, Thyroid, Chest X-ray (Text + Images + DOCX with images and context)

import streamlit as st
from PIL import Image
import io
import os
import re
import warnings
warnings.filterwarnings('ignore')
from dotenv import load_dotenv
load_dotenv()

# ============================================
# HARDCODE YOUR API KEY HERE (Optional - remove if using .env)
# ============================================
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")

# ============================================
# IMPORT RAG SYSTEM
# ============================================
try:
    from rag_system import MultimodalRAGSystem
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False
    print("RAG system not available")

# ============================================
# PAGE CONFIGURATION
# ============================================

st.set_page_config(
    page_title="Medical Report Analyzer",
    page_icon=":hospital:",
    layout="wide"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 1.8rem;
        text-align: center;
        padding: 0.8rem;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 1rem;
    }
    .stButton > button {
        background-color: #667eea;
        color: white;
        border-radius: 5px;
        font-weight: bold;
        width: 100%;
    }
    .result-box {
        background-color: #f0f2f6;
        padding: 1.5rem;
        border-radius: 10px;
        font-size: 1rem;
        line-height: 1.6;
        color: #000000;
    }
    .risk-low {
        background-color: #d4edda;
        color: #155724;
        padding: 0.5rem;
        border-radius: 10px;
        border-left: 5px solid #28a745;
        margin-bottom: 0.5rem;
    }
    .risk-moderate {
        background-color: #fff3cd;
        color: #856404;
        padding: 0.5rem;
        border-radius: 10px;
        border-left: 5px solid #ffc107;
        margin-bottom: 0.5rem;
    }
    .risk-high {
        background-color: #f8d7da;
        color: #721c24;
        padding: 0.5rem;
        border-radius: 10px;
        border-left: 5px solid #dc3545;
        margin-bottom: 0.5rem;
    }
</style>
""", unsafe_allow_html=True)

# ============================================
# INITIALIZE RAG SYSTEM
# ============================================

@st.cache_resource
def init_rag():
    """Initialize RAG system"""
    if RAG_AVAILABLE:
        try:
            return MultimodalRAGSystem()
        except Exception as e:
            st.sidebar.warning(f"RAG not available: {e}")
            return None
    return None

# ============================================
# FILE PARSING FUNCTIONS
# ============================================

def extract_from_pdf(file_bytes):
    """Extract text from PDF file"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        return ""

def extract_from_docx(file_bytes):
    """Extract text and images from DOCX file"""
    text = ""
    images = []
    clinical_context = ""
    
    try:
        import docx
        from docx import Document
        
        doc = Document(io.BytesIO(file_bytes))
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        # Auto-detect clinical context from text
        # Look for common clinical indicators
        text_lower = text.lower()
        context_keywords = ['patient', 'present', 'history', 'symptom', 'cough', 
                           'fever', 'pain', 'shortness', 'breath', 'age', 'male', 
                           'female', 'complains', 'reports', 'examination', 'clinical']
        
        if any(keyword in text_lower for keyword in context_keywords):
            clinical_context = text
        
        # Extract images from DOCX
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                images.append(image_data)
        
        return text.strip(), images, clinical_context
        
    except Exception as e:
        return f"Error extracting DOCX: {e}", [], ""

def extract_from_txt(file_bytes):
    """Extract text from TXT file"""
    try:
        return file_bytes.decode('utf-8').strip()
    except:
        try:
            return file_bytes.decode('latin-1').strip()
        except:
            return ""

def parse_uploaded_file(file_bytes, filename):
    """Main function to parse any uploaded file"""
    ext = filename.split('.')[-1].lower()
    
    if ext == 'pdf':
        return extract_from_pdf(file_bytes), 'pdf', None, ""
    elif ext == 'docx':
        text, images, context = extract_from_docx(file_bytes)
        return text, 'docx', images, context
    elif ext == 'txt':
        return extract_from_txt(file_bytes), 'txt', None, ""
    elif ext in ['png', 'jpg', 'jpeg', 'bmp', 'tiff']:
        return "", 'image', file_bytes, ""
    else:
        return "", 'unknown', None, ""

# ============================================
# MEDICAL VALUE EXTRACTION
# ============================================

def extract_medical_values(text):
    """Extract key medical values from text"""
    t = text.lower()
    values = {}
    
    # CBC values
    h = re.search(r'hemoglobin[^\d]*(\d+\.?\d*)', t) or re.search(r'hgb[^\d]*(\d+\.?\d*)', t)
    if h: values["Hemoglobin"] = float(h.group(1))
    
    rbc = re.search(r'red blood cell[^\d]*(\d+\.?\d*)', t) or re.search(r'rbc[^\d]*(\d+\.?\d*)', t)
    if rbc: values["RBC"] = float(rbc.group(1))
    
    wbc = re.search(r'white blood cell[^\d]*(\d+\.?\d*)', t) or re.search(r'wbc[^\d]*(\d+\.?\d*)', t)
    if wbc: values["WBC"] = float(wbc.group(1))
    
    platelets = re.search(r'platelet[^\d]*(\d+\.?\d*)', t) or re.search(r'plt[^\d]*(\d+\.?\d*)', t)
    if platelets: values["Platelets"] = float(platelets.group(1))
    
    hematocrit = re.search(r'hematocrit[^\d]*(\d+\.?\d*)', t) or re.search(r'hct[^\d]*(\d+\.?\d*)', t)
    if hematocrit: values["Hematocrit"] = float(hematocrit.group(1))
    
    mcv = re.search(r'mcv[^\d]*(\d+\.?\d*)', t)
    if mcv: values["MCV"] = float(mcv.group(1))
    
    mch = re.search(r'mch[^\d]*(\d+\.?\d*)', t)
    if mch: values["MCH"] = float(mch.group(1))
    
    rdw = re.search(r'rdw[^\d]*(\d+\.?\d*)', t)
    if rdw: values["RDW"] = float(rdw.group(1))
    
    # Thyroid values
    tsh = re.search(r'tsh[^\d]*(\d+\.?\d*)', t)
    if tsh: values["TSH"] = float(tsh.group(1))
    
    t3 = re.search(r't3[^\d]*(\d+\.?\d*)', t)
    if t3: values["T3"] = float(t3.group(1))
    
    t4 = re.search(r't4[^\d]*(\d+\.?\d*)', t)
    if t4: values["T4"] = float(t4.group(1))
    
    return values

def calculate_risk_from_values(values):
    """Calculate risk level from extracted lab values (CBC and Thyroid)"""
    risk_score = 0
    
    # CBC - Hemoglobin
    if 'Hemoglobin' in values:
        hgb = values['Hemoglobin']
        if hgb < 7:
            risk_score += 4
        elif hgb < 8:
            risk_score += 3
        elif hgb < 10:
            risk_score += 2
        elif hgb < 12:
            risk_score += 1
    
    # CBC - WBC
    if 'WBC' in values:
        wbc = values['WBC']
        if wbc > 20:
            risk_score += 4
        elif wbc > 15:
            risk_score += 3
        elif wbc > 11:
            risk_score += 2
        elif wbc < 3:
            risk_score += 2
    
    # CBC - Platelets
    if 'Platelets' in values:
        plt = values['Platelets']
        if plt < 50:
            risk_score += 3
        elif plt < 100:
            risk_score += 2
        elif plt > 600:
            risk_score += 2
    
    # CBC - MCV
    if 'MCV' in values:
        mcv = values['MCV']
        if mcv > 100 or mcv < 80:
            risk_score += 1
    
    # Thyroid - TSH
    if 'TSH' in values:
        tsh = values['TSH']
        if tsh > 20:
            risk_score += 4
        elif tsh > 10:
            risk_score += 3
        elif tsh > 4.5:
            risk_score += 2
        elif tsh < 0.1:
            risk_score += 3
        elif tsh < 0.4:
            risk_score += 2
    
    # Thyroid - T4
    if 'T4' in values:
        t4 = values['T4']
        if t4 < 2:
            risk_score += 2
        elif t4 < 4.5:
            risk_score += 1
    
    # Risk determination
    if risk_score >= 4:
        return "HIGH RISK - Immediate medical attention"
    elif risk_score >= 3:
        return "HIGH RISK"
    elif risk_score >= 2:
        return "MODERATE RISK"
    elif risk_score >= 1:
        return "LOW TO MODERATE RISK"
    return "LOW RISK - Normal range"

def calculate_risk_from_symptoms(text):
    """Calculate risk level from clinical symptoms and radiology findings (Chest X-ray)"""
    text_lower = text.lower()
    risk_score = 0
    condition = None
    
    # ============================================
    # DETECT CONDITION
    # ============================================
    
    # Pneumonia detection
    if 'pneumonia' in text_lower or 'consolidation' in text_lower or 'air bronchogram' in text_lower:
        condition = 'pneumonia'
        risk_score += 5
        
        # Severity modifiers
        if 'fever' in text_lower:
            risk_score += 2
            if '102' in text_lower or 'high fever' in text_lower:
                risk_score += 1
        if 'shortness of breath' in text_lower or 'sob' in text_lower:
            risk_score += 2
        if 'productive cough' in text_lower or 'sputum' in text_lower:
            risk_score += 1
        if 'o2 saturation' in text_lower:
            risk_score += 2
            if '94%' in text_lower or '90%' in text_lower:
                risk_score += 1
        if 'hypoxia' in text_lower:
            risk_score += 2
        if 'tachycardia' in text_lower or 'hr 9' in text_lower:
            risk_score += 1
        if 'tachypnea' in text_lower or 'rr 2' in text_lower:
            risk_score += 1
        if 'elderly' in text_lower or 'older' in text_lower:
            risk_score += 1
        if 'smoking' in text_lower or 'smoker' in text_lower:
            risk_score += 1
    
    # Cardiomegaly detection
    elif 'cardiomegaly' in text_lower or 'enlarged heart' in text_lower or 'cardiac silhouette' in text_lower:
        condition = 'cardiomegaly'
        risk_score += 4
        
        # Severity modifiers
        if 'shortness of breath' in text_lower or 'sob' in text_lower:
            risk_score += 2
        if 'orthopnea' in text_lower:
            risk_score += 2
        if 'pedal edema' in text_lower or 'leg swelling' in text_lower:
            risk_score += 2
        if 'congestion' in text_lower or 'pulmonary edema' in text_lower:
            risk_score += 2
        if 'hypertension' in text_lower or 'high blood pressure' in text_lower:
            risk_score += 1
        if 'heart failure' in text_lower or 'chf' in text_lower:
            risk_score += 2
    
    # Pleural effusion detection
    elif 'effusion' in text_lower or 'pleural' in text_lower:
        condition = 'effusion'
        risk_score += 3
        
        # Severity modifiers
        if 'large' in text_lower:
            risk_score += 2
        if 'bilateral' in text_lower:
            risk_score += 2
        if 'shortness of breath' in text_lower:
            risk_score += 2
    
    # Normal finding
    elif 'normal' in text_lower or 'clear' in text_lower:
        condition = 'normal'
        risk_score = 0
    
    # ============================================
    # RISK LEVEL DETERMINATION
    # ============================================
    
    if condition == 'pneumonia':
        if risk_score >= 8:
            return "HIGH RISK - Severe Pneumonia", condition
        elif risk_score >= 5:
            return "HIGH RISK - Pneumonia", condition
        else:
            return "MODERATE RISK - Mild Pneumonia", condition
    
    elif condition == 'cardiomegaly':
        if risk_score >= 7:
            return "HIGH RISK - Severe Cardiomegaly with CHF", condition
        elif risk_score >= 5:
            return "MODERATE TO HIGH RISK - Cardiomegaly with symptoms", condition
        else:
            return "MODERATE RISK - Cardiomegaly", condition
    
    elif condition == 'effusion':
        if risk_score >= 6:
            return "HIGH RISK - Large effusion", condition
        elif risk_score >= 4:
            return "MODERATE RISK - Moderate effusion", condition
        else:
            return "LOW TO MODERATE RISK - Small effusion", condition
    
    elif condition == 'normal':
        return "LOW RISK - Normal study", condition
    
    else:
        if risk_score >= 5:
            return "HIGH RISK", condition
        elif risk_score >= 3:
            return "MODERATE RISK", condition
        elif risk_score >= 1:
            return "LOW TO MODERATE RISK", condition
        return "LOW RISK", condition

# ============================================
# IMAGE ANALYSIS FUNCTION (for X-ray images)
# ============================================

def analyze_xray_image(image_bytes, clinical_context):
    """Generate analysis for chest X-ray image using clinical context"""
    try:
        from groq import Groq
        client = Groq(api_key=GROQ_API_KEY)
        
        prompt = f"""You are a medical AI assistant. Analyze this chest X-ray based on the clinical context.

CLINICAL CONTEXT:
{clinical_context}

Please provide a brief clinical analysis including:
1. Possible findings based on the symptoms
2. Clinical significance
3. Recommended next steps

Keep response under 200 words. Be professional but concise.

ANALYSIS:"""
        
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=400
        )
        
        return response.choices[0].message.content
    except Exception as e:
        return f"Analysis error: {e}"

# ============================================
# GROQ LLM ANALYSIS FOR TEXT
# ============================================

def get_groq_analysis(report_text, is_xray=False, clinical_context=""):
    """Get concise analysis from Groq Llama 3"""
    
    try:
        from groq import Groq
    except ImportError:
        return "Warning: Groq package not installed. Run: pip install groq"
    
    api_key = GROQ_API_KEY
    
    if not api_key:
        return "Warning: GROQ_API_KEY not set in the file."
    
    try:
        client = Groq(api_key=api_key)
        
        if is_xray:
            prompt = f"""You are a medical AI assistant. Analyze this chest X-ray based on the clinical context.

CLINICAL CONTEXT:
{clinical_context}

Please provide a brief clinical analysis including:
1. Possible findings based on the symptoms
2. Clinical significance
3. Recommended next steps

Keep response under 150 words. Be professional but concise."""
        else:
            values = extract_medical_values(report_text)
            values_str = ", ".join([f"{k}: {v}" for k, v in values.items()]) if values else "No specific values extracted"
            
            prompt = f"""You are a medical AI assistant. Analyze this medical report concisely.

REPORT TEXT:
{report_text[:1500]}

EXTRACTED VALUES:
{values_str}

Please provide a brief clinical analysis including:
1. Key abnormal findings (if any)
2. Clinical significance
3. One clear recommendation

Keep response under 150 words. Be professional but concise."""

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=350
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        error_msg = str(e)
        if "401" in error_msg:
            return "Warning: Invalid Groq API Key. Please check your key."
        elif "429" in error_msg:
            return "Warning: Rate limit exceeded. Please try again in a few moments."
        else:
            return f"Warning: Groq Error: {error_msg}"

# ============================================
# MAIN APP
# ============================================

def main():
    st.markdown('<div class="main-header">Medical Report Analyzer</div>', unsafe_allow_html=True)
    
    # Initialize RAG
    rag = init_rag()
    
    # Sidebar
    with st.sidebar:
        st.header("System Status")
        
        if GROQ_API_KEY and GROQ_API_KEY.startswith("gsk_"):
            st.success("Groq API Key Configured")
        else:
            st.error("Groq API Key Missing")
        
        if rag:
            stats = rag.get_stats()
            st.success(f"RAG Ready | {stats['total_records']} records")
        else:
            st.info("RAG not available (Milvus not running)")
        
        st.divider()
        
        st.header("Sample Reports")
        if st.button("CBC - Severe Anemia"):
            st.session_state['input_text'] = """CBC REPORT:
Hemoglobin: 6.5 g/dL (Low)
RBC: 1.8 M/mcL (Low)
Hematocrit: 19.5% (Low)
MCV: 109.6 fL (High)
MCH: 36.5 pg (High)
RDW: 16.0% (High)
WBC: 18.5 x10^3/uL (High)
Patient presents with fatigue and weakness."""
        
        if st.button("Thyroid - Hypothyroidism"):
            st.session_state['input_text'] = """THYROID FUNCTION:
TSH: 12.5 mIU/L (High)
T4: 3.2 mcg/dL (Low)
Patient has weight gain and fatigue."""
        
        if st.button("Chest X-ray - Pneumonia"):
            st.session_state['input_text'] = """CHEST X-RAY

CLINICAL CONTEXT:
65-year-old male with productive cough, fever (102°F), shortness of breath for 4 days. Smoking history.

FINDINGS:
Right lower lobe consolidation with air bronchograms.

IMPRESSION:
Community acquired pneumonia."""
        
        if st.button("Chest X-ray - Cardiomegaly"):
            st.session_state['input_text'] = """CHEST X-RAY

CLINICAL CONTEXT:
72-year-old female with shortness of breath, orthopnea, leg swelling. History of heart failure.

FINDINGS:
Cardiomegaly with pulmonary vascular congestion. Small bilateral pleural effusions.

IMPRESSION:
Congestive heart failure with cardiomegaly."""
        
        st.divider()
        
        st.markdown("### Supported Formats")
        st.markdown("""
        - Text: Direct paste (CBC, Thyroid, Radiology reports)
        - PDF/DOCX/TXT: Extract text from documents
        - X-ray Images: Upload image + provide clinical context
        """)
    
    # Main content area
    st.subheader("Upload or Paste Medical Report")
    
    input_type = st.radio("Select input method:", ["Text Input", "File Upload"], horizontal=True)
    
    report_text = ""
    is_xray = False
    clinical_context = ""
    uploaded_image = None
    
    if input_type == "Text Input":
        report_text = st.text_area(
            "Paste your medical report here:",
            height=250,
            placeholder="Paste CBC, Thyroid, or Chest X-ray report...",
            value=st.session_state.get('input_text', '')
        )
    else:
        uploaded_file = st.file_uploader(
            "Upload medical report (PDF, DOCX, TXT, PNG, JPG, JPEG):",
            type=['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'],
            help="Upload documents or X-ray images"
        )
        
        if uploaded_file:
            file_bytes = uploaded_file.read()
            extracted_text, file_type, extracted_images, extracted_context = parse_uploaded_file(file_bytes, uploaded_file.name)
            
            if file_type == 'docx':
                # Handle DOCX with possible images
                if extracted_text:
                    report_text = extracted_text
                    st.success(f"Successfully processed: {uploaded_file.name}")
                    with st.expander("View extracted text"):
                        st.text(report_text[:1000] + ("..." if len(report_text) > 1000 else ""))
                
                if extracted_images:
                    st.info(f"Found {len(extracted_images)} image(s) in the DOCX file")
                    uploaded_image = extracted_images[0]
                    st.image(uploaded_image, caption="Extracted Image from DOCX", width=350)
                    is_xray = True
                    
                    # Use extracted text as clinical context automatically
                    if extracted_context and len(extracted_context) > 50:
                        clinical_context = extracted_context
                        st.text_area("Clinical Context (auto-extracted from document):", value=clinical_context[:500], height=120, disabled=True)
                        st.success("Clinical context automatically extracted from document!")
                    else:
                        clinical_context = st.text_area(
                            "Clinical Context (required for X-ray analysis):",
                            height=120,
                            placeholder="Example: Patient is 65-year-old male with cough, fever (102F), shortness of breath for 3 days."
                        )
                        if clinical_context:
                            st.success("Clinical context provided")
                        else:
                            st.warning("Please provide clinical context for X-ray analysis")
                else:
                    # No images found, treat as text report
                    if not report_text:
                        st.error("Could not extract text from DOCX file.")
            
            elif file_type == 'image':
                is_xray = True
                uploaded_image = extracted_images
                st.image(uploaded_image, caption="Uploaded Chest X-ray", width=350)
                st.info("Chest X-ray detected. Please provide clinical context below.")
                clinical_context = st.text_area(
                    "Clinical Context (required for X-ray analysis):",
                    height=120,
                    placeholder="Example: Patient is 65-year-old male with cough, fever (102F), shortness of breath for 3 days."
                )
                if clinical_context:
                    st.success("Clinical context provided")
                else:
                    st.warning("Please provide clinical context for X-ray analysis")
            
            elif file_type in ['pdf', 'txt']:
                if extracted_text:
                    report_text = extracted_text
                    st.success(f"Successfully processed: {uploaded_file.name}")
                    with st.expander("View extracted text"):
                        st.text(report_text[:1000] + ("..." if len(report_text) > 1000 else ""))
                else:
                    st.error("Could not extract text from file.")
            
            else:
                st.error(f"Unsupported file type or could not process: {uploaded_file.name}")
    
    # Analyze button
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        analyze_btn = st.button("Analyze Report", type="primary", use_container_width=True)
    
    # Results
    if analyze_btn:
        if input_type == "Text Input" and not report_text:
            st.error("Please provide report text")
        elif input_type == "File Upload" and is_xray and not clinical_context and not report_text:
            st.error("Please provide clinical context for the X-ray image")
        elif input_type == "File Upload" and not is_xray and not report_text:
            st.error("Please upload a valid file")
        else:
            with st.spinner("Analyzing with Groq Llama 3..."):
                # Get risk level based on input type and content
                if not is_xray:
                    # Check if this is actually a chest X-ray report (even in text mode)
                    text_lower = report_text.lower()
                    is_chest_xray_text = any(word in text_lower for word in 
                        ['chest x-ray', 'chest radiograph', 'cxr', 'consolidation', 
                         'air bronchogram', 'pneumonia', 'lung', 'effusion', 'cardiomegaly'])
                    
                    if is_chest_xray_text:
                        # This is a chest X-ray report in text mode
                        risk_level, condition = calculate_risk_from_symptoms(report_text)
                        # Override risk for pneumonia
                        if 'pneumonia' in text_lower or 'consolidation' in text_lower:
                            risk_level = "HIGH RISK - Pneumonia detected"
                        elif 'cardiomegaly' in text_lower:
                            risk_level = "MODERATE TO HIGH RISK - Cardiomegaly"
                        elif 'effusion' in text_lower:
                            if 'HIGH' not in risk_level:
                                risk_level = "MODERATE RISK - Pleural effusion"
                    else:
                        # CBC or Thyroid report
                        values = extract_medical_values(report_text)
                        risk_level = calculate_risk_from_values(values)
                else:
                    # For chest X-ray image upload or DOCX with image
                    # Use the clinical context that was auto-extracted or manually entered
                    full_text = (clinical_context + " " + report_text).lower()
                    risk_level, condition = calculate_risk_from_symptoms(full_text)
                    
                    # Additional check for keywords
                    if 'pneumonia' in full_text or 'consolidation' in full_text:
                        if 'HIGH' not in risk_level:
                            risk_level = "HIGH RISK - Pneumonia detected"
                    elif 'cardiomegaly' in full_text:
                        if 'HIGH' not in risk_level:
                            risk_level = "MODERATE TO HIGH RISK - Cardiomegaly"
                    elif 'effusion' in full_text:
                        if 'HIGH' not in risk_level and 'MODERATE' not in risk_level:
                            risk_level = "MODERATE RISK - Pleural effusion"
                
                # Display risk level with color
                if "HIGH" in risk_level:
                    st.markdown(f'<div class="risk-high">Risk Level: {risk_level}</div>', unsafe_allow_html=True)
                elif "MODERATE" in risk_level:
                    st.markdown(f'<div class="risk-moderate">Risk Level: {risk_level}</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="risk-low">Risk Level: {risk_level}</div>', unsafe_allow_html=True)
                
                # RAG Similar Cases (only for text, not for images)
                if rag and not is_xray and report_text:
                    with st.spinner("Retrieving similar cases..."):
                        try:
                            similar_cases = rag.search_by_text(report_text, top_k=3)
                            if similar_cases:
                                with st.expander("Similar Cases from Database (RAG)", expanded=False):
                                    for i, case in enumerate(similar_cases):
                                        st.markdown(f"**Case {i+1}** - {case['report_type']} (Match: {case['score']:.2%})")
                                        st.markdown(f"*Risk: {case.get('risk_level', 'Unknown')} | Consult: {case.get('consultation', 'Monitor')}*")
                                        st.caption(case.get('narrative', '')[:200] + "...")
                                        st.divider()
                        except Exception as e:
                            st.warning(f"RAG search failed: {e}")
                
                # Doctor consultation recommendation
                if "HIGH" in risk_level:
                    st.warning("Consult Doctor Required - Immediate medical attention recommended")
                elif "MODERATE" in risk_level:
                    st.warning("Consult Doctor Required - Schedule appointment within 1-2 weeks")
                else:
                    st.info("Monitor at Home - Routine follow-up as clinically indicated")
                
                # Get Groq analysis - use appropriate method based on input type
                if is_xray:
                    # Use image analysis with clinical context
                    if clinical_context:
                        result = analyze_xray_image(None, clinical_context)
                    else:
                        result = "No clinical context provided. Please provide patient symptoms and history for analysis."
                else:
                    result = get_groq_analysis(report_text, is_xray=False)
                
                st.markdown("---")
                st.subheader("Analysis Result")
                st.markdown(f'<div class="result-box" style="color: #000000 !important;">{result}</div>', unsafe_allow_html=True)
                st.markdown("---")
                st.caption("Powered by Groq Llama 3 | For informational purposes only. Consult a healthcare professional for medical advice.")

if __name__ == "__main__":
    main()