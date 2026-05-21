# parser.py
# Medical Report Parser - Handles PDF, DOCX, TXT, and Image files

import os
import re
import io
from PIL import Image
import pytesseract
from PyPDF2 import PdfReader
import docx
from typing import Dict, List, Tuple, Optional
import warnings
warnings.filterwarnings('ignore')

class MedicalReportParser:
    """Parse medical reports from PDF, DOCX, TXT, and Image files"""
    
    def __init__(self, tesseract_path: Optional[str] = None):
        """
        Initialize parser
        Args:
            tesseract_path: Path to tesseract executable (Windows only)
        """
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        print("Medical Report Parser initialized")
    
    # ============================================
    # TEXT EXTRACTION FROM FILES
    # ============================================
    
    def extract_from_pdf(self, file_content: bytes) -> Tuple[str, str]:
        """
        Extract text from PDF file
        Returns: (extracted_text, file_type)
        """
        try:
            pdf = PdfReader(io.BytesIO(file_content))
            text = ""
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page_text + "\n"
            
            if text.strip():
                return text.strip(), 'pdf'
            else:
                return "No text found in PDF (may be scanned image PDF)", 'pdf'
        except Exception as e:
            return f"Error extracting PDF: {str(e)}", 'pdf'
    
    def extract_from_docx(self, file_content: bytes) -> Tuple[str, str]:
        """
        Extract text from DOCX file
        Returns: (extracted_text, file_type)
        """
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if text.strip():
                return text.strip(), 'docx'
            else:
                return "No text found in DOCX file", 'docx'
        except Exception as e:
            return f"Error extracting DOCX: {str(e)}", 'docx'
    
    def extract_from_txt(self, file_content: bytes) -> Tuple[str, str]:
        """
        Extract text from TXT file
        Returns: (extracted_text, file_type)
        """
        try:
            # Try UTF-8 first
            text = file_content.decode('utf-8').strip()
            return text, 'txt'
        except UnicodeDecodeError:
            try:
                # Try Latin-1 as fallback
                text = file_content.decode('latin-1').strip()
                return text, 'txt'
            except Exception as e:
                return f"Error decoding text file: {str(e)}", 'txt'
    
    def extract_from_image(self, file_content: bytes) -> Tuple[str, str]:
        """
        Extract text from image using OCR
        Returns: (extracted_text, file_type)
        """
        try:
            image = Image.open(io.BytesIO(file_content))
            
            # Preprocess image for better OCR
            image = image.convert('L')  # Convert to grayscale
            
            # Apply threshold to make text clearer
            threshold = 150
            image = image.point(lambda p: p > threshold and 255)
            
            # OCR
            text = pytesseract.image_to_string(image)
            
            if text.strip():
                return text.strip(), 'image'
            else:
                return "No text detected in image", 'image'
        except Exception as e:
            return f"Error extracting text from image: {str(e)}", 'image'
    
    def parse_file(self, file_content: bytes, filename: str) -> Tuple[str, str]:
        """
        Main method to parse any uploaded file
        Returns: (extracted_text, file_type)
        """
        file_ext = filename.split('.')[-1].lower()
        
        if file_ext == 'pdf':
            return self.extract_from_pdf(file_content)
        elif file_ext == 'docx':
            return self.extract_from_docx(file_content)
        elif file_ext == 'txt':
            return self.extract_from_txt(file_content)
        elif file_ext in ['png', 'jpg', 'jpeg', 'tiff', 'bmp', 'gif']:
            return self.extract_from_image(file_content)
        else:
            return f"Unsupported file type: {file_ext}", 'unknown'
    
    # ============================================
    # REPORT TYPE DETECTION
    # ============================================
    
    def detect_report_type(self, text: str) -> str:
        """Detect the type of medical report from text"""
        text_lower = text.lower()
        
        # CBC keywords
        cbc_keywords = [
            'hemoglobin', 'hgb', 'hematocrit', 'hct', 'wbc', 'white blood cell',
            'rbc', 'red blood cell', 'platelet', 'plt', 'cbc', 'complete blood count',
            'mcv', 'mch', 'mchc', 'rdw'
        ]
        
        # Thyroid keywords
        thyroid_keywords = [
            'tsh', 't3', 't4', 'thyroxine', 'triiodothyronine', 'thyroid',
            'ft3', 'ft4', 'thyroid function', 'tft', 'free t3', 'free t4'
        ]
        
        # Chest X-ray keywords
        xray_keywords = [
            'chest x-ray', 'chest radiograph', 'cxr', 'x-ray chest', 'lung',
            'pneumonia', 'consolidation', 'infiltrate', 'effusion', 'pleural',
            'cardiomegaly', 'pulmonary', 'bronchus', 'alveolar', 'interstitial'
        ]
        
        # Calculate scores
        scores = {
            'CBC': sum(1 for kw in cbc_keywords if kw in text_lower),
            'Thyroid': sum(1 for kw in thyroid_keywords if kw in text_lower),
            'Chest X-ray': sum(1 for kw in xray_keywords if kw in text_lower)
        }
        
        # Get the highest score
        detected = max(scores, key=scores.get)
        
        if scores[detected] == 0:
            return "General Medical Report"
        return detected
    
    # ============================================
    # CBC PARSING
    # ============================================
    
    def parse_cbc(self, text: str) -> Dict:
        """Parse CBC report values"""
        text_lower = text.lower()
        values = {}
        
        # Patterns for extracting values
        patterns = {
            'Hemoglobin': [
                r'hemoglobin:?\s*(\d+\.?\d*)',
                r'hgb:?\s*(\d+\.?\d*)',
                r'hb:?\s*(\d+\.?\d*)'
            ],
            'WBC': [
                r'wbc:?\s*(\d+\.?\d*)',
                r'white blood cell:?\s*(\d+\.?\d*)',
                r'leukocyte:?\s*(\d+\.?\d*)'
            ],
            'Platelets': [
                r'platelet:?\s*(\d+\.?\d*)',
                r'plt:?\s*(\d+\.?\d*)',
                r'thrombocyte:?\s*(\d+\.?\d*)'
            ],
            'RBC': [
                r'rbc:?\s*(\d+\.?\d*)',
                r'red blood cell:?\s*(\d+\.?\d*)',
                r'erythrocyte:?\s*(\d+\.?\d*)'
            ],
            'Hematocrit': [
                r'hematocrit:?\s*(\d+\.?\d*)',
                r'hct:?\s*(\d+\.?\d*)'
            ],
            'MCV': [
                r'mcv:?\s*(\d+\.?\d*)',
                r'mean corpuscular volume:?\s*(\d+\.?\d*)'
            ],
            'MCH': [
                r'mch:?\s*(\d+\.?\d*)',
                r'mean corpuscular hemoglobin:?\s*(\d+\.?\d*)'
            ]
        }
        
        for key, pattern_list in patterns.items():
            for pattern in pattern_list:
                match = re.search(pattern, text_lower)
                if match:
                    values[key] = float(match.group(1))
                    break
        
        # Calculate risk level
        risk = self._calculate_cbc_risk(values)
        
        # Generate interpretation
        interpretation = self._generate_cbc_interpretation(values)
        
        return {
            'type': 'CBC',
            'values': values,
            'interpretation': interpretation,
            'risk_level': risk
        }
    
    def _calculate_cbc_risk(self, values: Dict) -> str:
        """Calculate risk level from CBC values"""
        risk_score = 0
        
        if 'Hemoglobin' in values:
            hgb = values['Hemoglobin']
            if hgb < 10:
                risk_score += 3
            elif hgb < 12:
                risk_score += 2
            elif hgb > 18:
                risk_score += 2
        
        if 'WBC' in values:
            wbc = values['WBC']
            if wbc > 15:
                risk_score += 3
            elif wbc > 11:
                risk_score += 2
            elif wbc < 3:
                risk_score += 2
        
        if risk_score >= 5:
            return "HIGH RISK"
        elif risk_score >= 3:
            return "MODERATE RISK"
        elif risk_score >= 1:
            return "LOW TO MODERATE RISK"
        return "LOW RISK"
    
    def _generate_cbc_interpretation(self, values: Dict) -> List[str]:
        """Generate clinical interpretation for CBC"""
        interpretation = []
        
        if 'Hemoglobin' in values:
            hgb = values['Hemoglobin']
            if hgb < 12:
                interpretation.append(f"Low hemoglobin ({hgb} g/dL) indicates anemia")
            elif hgb > 17:
                interpretation.append(f"High hemoglobin ({hgb} g/dL) suggests polycythemia")
            else:
                interpretation.append(f"Hemoglobin is normal ({hgb} g/dL)")
        
        if 'WBC' in values:
            wbc = values['WBC']
            if wbc > 11:
                interpretation.append(f"Elevated WBC ({wbc}) suggests infection or inflammation")
            elif wbc < 4:
                interpretation.append(f"Low WBC ({wbc}) indicates leukopenia")
            else:
                interpretation.append(f"WBC is normal ({wbc})")
        
        if 'Platelets' in values:
            plt = values['Platelets']
            if plt > 450:
                interpretation.append(f"High platelets ({plt}) indicates thrombocytosis")
            elif plt < 150:
                interpretation.append(f"Low platelets ({plt}) indicates thrombocytopenia")
            else:
                interpretation.append(f"Platelets are normal ({plt})")
        
        if not interpretation:
            interpretation.append("All CBC values are within normal range")
        
        return interpretation
    
    # ============================================
    # THYROID PARSING
    # ============================================
    
    def parse_thyroid(self, text: str) -> Dict:
        """Parse Thyroid report values"""
        text_lower = text.lower()
        values = {}
        
        patterns = {
            'TSH': [r'tsh:?\s*(\d+\.?\d*)'],
            'T3': [r't3:?\s*(\d+\.?\d*)'],
            'T4': [r't4:?\s*(\d+\.?\d*)', r'thyroxine:?\s*(\d+\.?\d*)'],
            'Free T4': [r'free t4:?\s*(\d+\.?\d*)', r'ft4:?\s*(\d+\.?\d*)']
        }
        
        for key, pattern_list in patterns.items():
            for pattern in pattern_list:
                match = re.search(pattern, text_lower)
                if match:
                    values[key] = float(match.group(1))
                    break
        
        risk = self._calculate_thyroid_risk(values)
        interpretation = self._generate_thyroid_interpretation(values)
        
        return {
            'type': 'Thyroid',
            'values': values,
            'interpretation': interpretation,
            'risk_level': risk
        }
    
    def _calculate_thyroid_risk(self, values: Dict) -> str:
        """Calculate risk level from Thyroid values"""
        if 'TSH' in values:
            tsh = values['TSH']
            if tsh > 10:
                return "HIGH RISK"
            elif tsh > 4.5:
                return "MODERATE RISK"
            elif tsh < 0.1:
                return "HIGH RISK"
            elif tsh < 0.4:
                return "MODERATE RISK"
        return "LOW RISK"
    
    def _generate_thyroid_interpretation(self, values: Dict) -> List[str]:
        """Generate clinical interpretation for Thyroid"""
        interpretation = []
        
        if 'TSH' in values:
            tsh = values['TSH']
            if tsh > 4.5:
                interpretation.append(f"Elevated TSH ({tsh}) suggests HYPOTHYROIDISM")
            elif tsh < 0.4:
                interpretation.append(f"Low TSH ({tsh}) suggests HYPERTHYROIDISM")
            else:
                interpretation.append(f"TSH is normal ({tsh})")
        
        if not interpretation:
            interpretation.append("Thyroid function appears normal")
        
        return interpretation
    
    # ============================================
    # CHEST X-RAY PARSING
    # ============================================
    
    def parse_chest_xray(self, text: str) -> Dict:
        """Parse Chest X-ray report"""
        text_lower = text.lower()
        
        findings = []
        
        # Find key findings
        finding_keywords = {
            'Pneumonia/Consolidation': ['pneumonia', 'consolidation', 'infiltrate', 'air bronchogram'],
            'Pleural effusion': ['effusion', 'pleural', 'costophrenic angle blunting'],
            'Cardiomegaly': ['cardiomegaly', 'enlarged heart', 'cardiac silhouette'],
            'Pulmonary edema': ['edema', 'pulmonary congestion', 'vascular congestion'],
            'Normal': ['normal', 'no abnormality', 'clear lungs', 'unremarkable']
        }
        
        for finding, keywords in finding_keywords.items():
            if any(kw in text_lower for kw in keywords):
                findings.append(finding)
        
        # Determine risk level
        if 'pneumonia' in text_lower or 'consolidation' in text_lower:
            risk = "HIGH RISK"
        elif 'effusion' in text_lower:
            risk = "MODERATE RISK"
        elif 'cardiomegaly' in text_lower:
            risk = "MODERATE TO HIGH RISK"
        else:
            risk = "LOW RISK"
        
        interpretation = f"Findings: {', '.join(findings) if findings else 'No significant findings detected'}"
        
        return {
            'type': 'Chest X-ray',
            'findings': findings,
            'interpretation': [interpretation],
            'risk_level': risk
        }
    
    # ============================================
    # MAIN PARSE METHOD
    # ============================================
    
    def parse(self, text: str) -> Dict:
        """
        Main method to parse any medical report text
        Returns structured data with type, values, interpretation, and risk
        """
        if not text or len(text.strip()) < 10:
            return {
                'type': 'Unknown',
                'values': {},
                'interpretation': ['Insufficient text to analyze'],
                'risk_level': 'UNKNOWN'
            }
        
        report_type = self.detect_report_type(text)
        
        if report_type == 'CBC':
            return self.parse_cbc(text)
        elif report_type == 'Thyroid':
            return self.parse_thyroid(text)
        elif report_type == 'Chest X-ray':
            return self.parse_chest_xray(text)
        else:
            return {
                'type': report_type,
                'values': {},
                'interpretation': ['Unable to determine report type. Please ensure the report is CBC, Thyroid, or Chest X-ray.'],
                'risk_level': 'UNKNOWN'
            }


# ============================================
# TEST FUNCTION
# ============================================

if __name__ == "__main__":
    parser = MedicalReportParser()
    
    # Test CBC
    cbc_text = """CBC REPORT:
    Patient: John Doe
    Hemoglobin: 10.2 g/dL
    WBC: 14.5 x10^3/uL
    Platelets: 180 x10^3/uL
    RBC: 3.8 x10^6/uL
    Hematocrit: 32%"""
    
    print("\n" + "="*60)
    print("TESTING CBC PARSER")
    print("="*60)
    result = parser.parse(cbc_text)
    print(f"Type: {result['type']}")
    print(f"Values: {result['values']}")
    print(f"Interpretation: {result['interpretation']}")
    print(f"Risk: {result['risk_level']}")
    
    # Test Thyroid
    thyroid_text = """THYROID FUNCTION:
    TSH: 12.5 mIU/L
    T4: 3.2 mcg/dL
    Patient has weight gain and fatigue."""
    
    print("\n" + "="*60)
    print("TESTING THYROID PARSER")
    print("="*60)
    result = parser.parse(thyroid_text)
    print(f"Type: {result['type']}")
    print(f"Values: {result['values']}")
    print(f"Interpretation: {result['interpretation']}")
    print(f"Risk: {result['risk_level']}")
    
    # Test Chest X-ray
    xray_text = """CHEST X-RAY:
    FINDINGS: Right lower lobe consolidation with air bronchograms.
    IMPRESSION: Community acquired pneumonia."""
    
    print("\n" + "="*60)
    print("TESTING CHEST X-RAY PARSER")
    print("="*60)
    result = parser.parse(xray_text)
    print(f"Type: {result['type']}")
    print(f"Findings: {result.get('findings', [])}")
    print(f"Interpretation: {result['interpretation']}")
    print(f"Risk: {result['risk_level']}")